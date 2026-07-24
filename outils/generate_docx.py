#!/usr/bin/env python3
"""Convert 6 HTML blog articles to Word .docx files using python-docx."""

import os
import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt

BASE_DIR = "/Users/julianetaulelle/Documents/eliott"
OUT_DIR = os.path.join(BASE_DIR, "export-word")

ARTICLES = [
    "entretien-chaudiere-hautes-alpes.html",
    "adoucisseur-eau-hautes-alpes.html",
    "pose-poele-granules-hautes-alpes.html",
    "fuite-eau-urgence-hautes-alpes.html",
    "chauffe-eau-en-panne-hautes-alpes.html",
    "pression-chaudiere-hautes-alpes.html",
]


def strip_tags(html_text):
    """Remove all HTML tags and decode basic entities, returning plain text."""
    # Replace <br> variants with space
    html_text = re.sub(r'<br\s*/?>', ' ', html_text, flags=re.IGNORECASE)
    # Remove all other tags
    html_text = re.sub(r'<[^>]+>', '', html_text)
    # Decode common entities
    html_text = html_text.replace('&nbsp;', ' ')
    html_text = html_text.replace('&amp;', '&')
    html_text = html_text.replace('&lt;', '<')
    html_text = html_text.replace('&gt;', '>')
    html_text = html_text.replace('&quot;', '"')
    html_text = html_text.replace('&#39;', "'")
    html_text = html_text.replace('&eacute;', 'é')
    html_text = html_text.replace('&egrave;', 'è')
    html_text = html_text.replace('&ecirc;', 'ê')
    html_text = html_text.replace('&agrave;', 'à')
    html_text = html_text.replace('&acirc;', 'â')
    html_text = html_text.replace('&ocirc;', 'ô')
    html_text = html_text.replace('&ucirc;', 'û')
    html_text = html_text.replace('&ccedil;', 'ç')
    # Collapse whitespace
    html_text = re.sub(r'\s+', ' ', html_text).strip()
    return html_text


class ArticleParser(HTMLParser):
    """Parse the article-content div and extract structured content."""

    def __init__(self):
        super().__init__()
        self.in_article = False
        self.depth = 0          # nesting depth inside article-content
        self.current_tag = None
        self.current_attrs = {}
        self.current_box = None  # 'highlight' | 'warning' | None
        self.in_box = False
        self.content = []       # list of (type, text)
        # type values: 'h1','h2','h3','p','li','highlight-box','warning-box'
        self._buf = ''
        self._stack = []        # tag stack inside article
        self._li_buf = ''
        self._in_li = False

    def _attrs_dict(self, attrs):
        return {k: v for k, v in attrs}

    def handle_starttag(self, tag, attrs):
        ad = self._attrs_dict(attrs)
        classes = ad.get('class', '')

        # Detect entry into article-content
        if tag == 'div' and 'article-content' in classes:
            self.in_article = True
            self.depth = 1
            self._stack.append(('div', classes))
            return

        if not self.in_article:
            return

        # Track div depth
        if tag == 'div':
            self.depth += 1
            self._stack.append(('div', classes))
            # Detect special boxes
            if 'highlight-box' in classes:
                self.current_box = 'highlight'
                self.in_box = True
            elif 'warning-box' in classes:
                self.current_box = 'warning'
                self.in_box = True
            return

        if tag in ('h1', 'h2', 'h3'):
            self.current_tag = tag
            self._buf = ''
        elif tag == 'p':
            self.current_tag = 'p'
            self._buf = ''
        elif tag == 'li':
            self._in_li = True
            self._li_buf = ''
        elif tag == 'ul':
            pass  # nothing special on ul open
        elif tag in ('strong', 'em', 'b', 'i', 'a', 'span', 'sup', 'sub'):
            pass  # inline tags — text will be captured via handle_data

    def handle_endtag(self, tag):
        if not self.in_article:
            return

        if tag == 'div':
            if self._stack:
                popped = self._stack.pop()
                popped_class = popped[1] if len(popped) > 1 else ''
                if 'highlight-box' in popped_class or 'warning-box' in popped_class:
                    self.in_box = False
                    self.current_box = None
            self.depth -= 1
            if self.depth == 0:
                self.in_article = False
            return

        if tag in ('h1', 'h2', 'h3'):
            text = strip_tags(self._buf).strip()
            if text:
                self.content.append((tag, text))
            self.current_tag = None
            self._buf = ''
        elif tag == 'p':
            text = strip_tags(self._buf).strip()
            if text:
                if self.in_box and self.current_box == 'highlight':
                    self.content.append(('highlight-box', text))
                elif self.in_box and self.current_box == 'warning':
                    self.content.append(('warning-box', text))
                else:
                    self.content.append(('p', text))
            self.current_tag = None
            self._buf = ''
        elif tag == 'li':
            text = strip_tags(self._li_buf).strip()
            if text:
                self.content.append(('li', text))
            self._in_li = False
            self._li_buf = ''

    def handle_data(self, data):
        if not self.in_article:
            return
        if self._in_li:
            self._li_buf += data
        elif self.current_tag in ('h1', 'h2', 'h3', 'p'):
            self._buf += data

    def handle_entityref(self, name):
        # Older-style entity refs (rare in html5 but just in case)
        self.handle_data(f'&{name};')

    def handle_charref(self, name):
        self.handle_data(f'&#{name};')


def extract_article_content(html_path):
    """Return a list of (type, text) tuples from the article-content div."""
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()
    parser = ArticleParser()
    parser.feed(html)
    return parser.content


def build_docx(content, out_path):
    """Build a .docx from the extracted content list."""
    doc = Document()

    for item_type, text in content:
        if item_type == 'h1':
            doc.add_heading(text, level=1)
        elif item_type == 'h2':
            doc.add_heading(text, level=2)
        elif item_type == 'h3':
            doc.add_heading(text, level=3)
        elif item_type == 'p':
            doc.add_paragraph(text)
        elif item_type == 'li':
            doc.add_paragraph(text, style='List Bullet')
        elif item_type == 'highlight-box':
            doc.add_paragraph('ℹ️  ' + text)
        elif item_type == 'warning-box':
            doc.add_paragraph('⚠️  ' + text)

    doc.save(out_path)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    for filename in ARTICLES:
        html_path = os.path.join(BASE_DIR, filename)
        stem = os.path.splitext(filename)[0]
        out_path = os.path.join(OUT_DIR, stem + '.docx')

        print(f"Processing: {filename}")
        content = extract_article_content(html_path)
        print(f"  → {len(content)} elements extracted")

        # Also grab h1 from the article-hero section if not already present
        # (some files put h1 outside article-content)
        has_h1 = any(t == 'h1' for t, _ in content)
        if not has_h1:
            # Read the full file and find the first <h1>
            with open(html_path, 'r', encoding='utf-8') as f:
                raw = f.read()
            m = re.search(r'<h1[^>]*>(.*?)</h1>', raw, re.DOTALL | re.IGNORECASE)
            if m:
                h1_text = strip_tags(m.group(1)).strip()
                content.insert(0, ('h1', h1_text))
                print(f"  → h1 injected from hero: {h1_text[:60]}")

        build_docx(content, out_path)
        print(f"  → Saved: {out_path}")

    print("\nDone. All files generated in:", OUT_DIR)


if __name__ == '__main__':
    main()
